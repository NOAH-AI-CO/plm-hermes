# -*- coding: utf-8 -*-
"""
Tests for AttachmentDownloader

Run with: python -m utils.web_search.test_attachment_downloader

Note: Some tests require Azure Blob Storage configuration.
Tests marked with [INTEGRATION] require network access.
"""

import io
import asyncio
import logging
from unittest.mock import Mock, AsyncMock, patch

from utils.web_search.attachment_downloader import AttachmentDownloader, ParsedContent, DownloadResult
from utils.web_search.attachment_detector import AttachmentType, DetectedAttachment

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MockBlobStorageClient:
    """Mock blob storage client for testing"""

    def __init__(self):
        self._storage = {}  # blob_path -> (content, metadata)

    def upload_file(self, container, blob, file_obj, metadata=None):
        content = file_obj.read()
        self._storage[blob] = (content, metadata or {})
        return True

    def load_file(self, container, blob):
        if blob in self._storage:
            return self._storage[blob][0]
        return None

    def get_blob_meta(self, container, blob):
        if blob in self._storage:
            return self._storage[blob][1]
        return None


async def test_parse_pdf():
    """Test PDF parsing"""
    print("\n=== Test: PDF Parsing ===")

    # Create a simple PDF using reportlab if available, otherwise skip
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        # Create PDF in memory
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 700, "This is a test content.")
        c.save()
        pdf_bytes = buffer.getvalue()

        downloader = AttachmentDownloader()
        result = await downloader._parse_pdf(pdf_bytes, "test.pdf")

        print(f"  Text preview length: {len(result.text_preview)}")
        print(f"  Data description: {result.data_description}")

        assert "test.pdf" in result.data_description.lower()
        print("[PASS] PDF parsing")

    except ImportError:
        print("[SKIP] reportlab not installed, skipping PDF test")


async def test_parse_csv():
    """Test CSV parsing"""
    print("\n=== Test: CSV Parsing ===")

    csv_content = b"""name,age,city
Alice,30,New York
Bob,25,San Francisco
Charlie,35,Chicago
"""

    downloader = AttachmentDownloader()
    result = await downloader._parse_csv(csv_content, "data.csv")

    print(f"  Text preview:\n{result.text_preview[:200]}")
    print(f"  Data description: {result.data_description}")

    assert "name" in result.text_preview
    assert "age" in result.text_preview
    assert "3 rows" in result.data_description
    print("[PASS] CSV parsing")


async def test_parse_excel():
    """Test Excel parsing"""
    print("\n=== Test: Excel Parsing ===")

    try:
        import pandas as pd
        from openpyxl import Workbook

        # Create Excel in memory
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Name'
        ws['B1'] = 'Value'
        ws['A2'] = 'Item1'
        ws['B2'] = 100
        ws['A3'] = 'Item2'
        ws['B3'] = 200

        buffer = io.BytesIO()
        wb.save(buffer)
        excel_bytes = buffer.getvalue()

        downloader = AttachmentDownloader()
        result = await downloader._parse_excel(excel_bytes, "data.xlsx")

        print(f"  Text preview:\n{result.text_preview[:200]}")
        print(f"  Data description: {result.data_description}")

        assert "Name" in result.text_preview or "name" in result.text_preview.lower()
        print("[PASS] Excel parsing")

    except ImportError:
        print("[SKIP] openpyxl not installed, skipping Excel test")


async def test_parse_word():
    """Test Word parsing"""
    print("\n=== Test: Word Parsing ===")

    try:
        from docx import Document

        # Create Word document in memory
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_paragraph('Another paragraph with more text.')

        buffer = io.BytesIO()
        doc.save(buffer)
        word_bytes = buffer.getvalue()

        downloader = AttachmentDownloader()
        result = await downloader._parse_word(word_bytes, "doc.docx")

        print(f"  Text preview:\n{result.text_preview[:200]}")
        print(f"  Data description: {result.data_description}")

        assert "test" in result.text_preview.lower()
        print("[PASS] Word parsing")

    except ImportError:
        print("[SKIP] python-docx not installed, skipping Word test")


async def test_extract_filename_from_url():
    """Test filename extraction from URL"""
    print("\n=== Test: Filename Extraction ===")

    downloader = AttachmentDownloader()

    test_cases = [
        ("https://example.com/reports/annual-2024.pdf", "annual-2024.pdf"),
        ("https://example.com/file.xlsx?token=abc", "file.xlsx"),
        ("https://example.com/data%20file.csv", "data file.csv"),
        ("https://example.com/", "unknown"),
    ]

    for url, expected in test_cases:
        result = downloader._extract_filename_from_url(url)
        print(f"  {url} -> {result}")
        if expected != "unknown":
            assert expected in result or result == expected, f"Expected {expected}, got {result}"

    print("[PASS] Filename extraction")


async def test_detect_type():
    """Test file type detection"""
    print("\n=== Test: Type Detection ===")

    downloader = AttachmentDownloader()

    test_cases = [
        ("report.pdf", AttachmentType.PDF),
        ("data.xlsx", AttachmentType.EXCEL),
        ("data.xls", AttachmentType.EXCEL),
        ("export.csv", AttachmentType.CSV),
        ("document.docx", AttachmentType.WORD),
        ("document.doc", AttachmentType.WORD),
        ("image.png", AttachmentType.IMAGE),
        ("photo.jpg", AttachmentType.IMAGE),
        ("photo.jpeg", AttachmentType.IMAGE),
        ("animation.gif", AttachmentType.IMAGE),
        ("icon.bmp", AttachmentType.IMAGE),
        ("banner.webp", AttachmentType.IMAGE),
    ]

    for filename, expected_type in test_cases:
        result = downloader._detect_type(filename)
        print(f"  {filename} -> {result.value}")
        assert result == expected_type, f"Expected {expected_type}, got {result}"

    print("[PASS] Type detection")


async def test_blob_storage_roundtrip():
    """Test saving and fetching from blob storage"""
    print("\n=== Test: Blob Storage Roundtrip ===")

    # Create downloader with mock blob storage
    downloader = AttachmentDownloader()
    mock_client = MockBlobStorageClient()
    downloader.blob_storage_client = mock_client
    downloader.blob_container = "test-container"

    # Test data
    test_content = b"This is test file content"
    test_filename = "test_file.txt"
    test_url = "https://example.com/test_file.txt"

    # Save to blob
    blob_path = downloader._save_attachment_to_blob(
        url=test_url,
        file_bytes=test_content,
        filename=test_filename
    )

    print(f"  Saved to blob_path: {blob_path}")
    assert blob_path is not None

    # Fetch from blob
    fetched_content, fetched_filename = downloader.fetch_attachment_from_blob(blob_path)

    print(f"  Fetched filename: {fetched_filename}")
    print(f"  Fetched content length: {len(fetched_content) if fetched_content else 0}")

    assert fetched_content == test_content
    assert fetched_filename == test_filename

    print("[PASS] Blob storage roundtrip")


async def test_download_single_mock():
    """Test download_single with mocked HTTP"""
    print("\n=== Test: Download Single (Mocked) ===")

    csv_content = b"col1,col2\nval1,val2\n"

    downloader = AttachmentDownloader()
    mock_client = MockBlobStorageClient()
    downloader.blob_storage_client = mock_client
    downloader.blob_container = "test-container"

    # Mock the HTTP download
    with patch.object(downloader, '_download_to_memory', new_callable=AsyncMock) as mock_download:
        mock_download.return_value = (csv_content, None)

        result = await downloader.download_single(
            url="https://example.com/data.csv",
            filename="data.csv"
        )

        print(f"  Success: {result.success}")
        print(f"  Filename: {result.filename}")
        print(f"  Blob path: {result.blob_path}")
        print(f"  Text preview: {result.text_preview[:100] if result.text_preview else 'None'}")

        assert result.success
        assert result.filename == "data.csv"
        assert result.blob_path is not None

    print("[PASS] Download single (mocked)")


async def run_all_tests():
    """Run all tests"""
    print("=" * 60)
    print("AttachmentDownloader Tests")
    print("=" * 60)

    tests = [
        test_extract_filename_from_url,
        test_detect_type,
        test_parse_csv,
        test_parse_pdf,
        test_parse_excel,
        test_parse_word,
        test_blob_storage_roundtrip,
        test_download_single_mock,
    ]

    passed = 0
    failed = 0
    skipped = 0

    for test in tests:
        try:
            await test()
            passed += 1
        except AssertionError as e:
            print(f"[FAIL] {test.__name__}: {e}")
            failed += 1
        except Exception as e:
            if "SKIP" in str(e) or "not installed" in str(e):
                skipped += 1
            else:
                print(f"[ERROR] {test.__name__}: {e}")
                failed += 1

    print("\n" + "=" * 60)
    print(f"Results: {passed} passed, {failed} failed, {skipped} skipped")
    print("=" * 60)

    return failed == 0


if __name__ == "__main__":
    success = asyncio.run(run_all_tests())
    exit(0 if success else 1)
