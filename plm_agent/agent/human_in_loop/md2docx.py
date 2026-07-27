from typing import Dict, List, Optional
import json
import hashlib
import tempfile
import shutil
from pathlib import Path
import pypandoc
from mermaid import Mermaid
import vl_convert as vlc
from docx import Document
from docx.shared import Pt, Inches

# Define tool class
class MarkdownToWordTool():
    def run(
        self, 
        input_path: str, 
        output_path: Optional[str] = None, 
        logo_path='static/logo.png'
    ) -> str:
        """Execute Markdown to Word conversion"""
        try:
            # Create converter instance
            converter = self.MarkdownConverter(input_path, output_path, logo_path=logo_path)

            # Execute conversion
            success = converter.convert()
            
            if success:
                return f"Conversion successful! Output file saved at: {converter.output_path}"
            else:
                return "Conversion failed, please check error messages."
        except Exception as e:
            return f"Error during conversion process: {str(e)}"
    
    class DiagramRenderer:
        """Handles rendering diagrams to image files"""
        
        def __init__(self, output_dir: Path):
            self.output_dir = output_dir
            
        def render(self, resources: List[Dict]) -> None:
            """Render all detected diagrams"""
            for resource in resources:
                try:
                    if resource['type'] == 'mermaid':
                        self._render_mermaid(resource)
                    elif resource['type'] in ['vega', 'vega-lite']:
                        self._render_vega(resource)
                except Exception as e:
                    print(f"Unable to render {resource['type']} diagram: {str(e)}")
                    resource['error'] = True
        
        def _render_mermaid(self, resource: Dict) -> None:
            """High-resolution rendering for complex diagrams"""
            output_file = self.output_dir / f"mermaid_{resource['hash']}.png"
            
            # For complex diagrams (e.g., with many nodes/text)
            diagram_script = resource['content']
            base_width = 1000  # Base width
            high_scale = 2.5   # Higher scale factor
            
            mermaid = Mermaid(
                graph=diagram_script,
                width=base_width,
                scale=high_scale
            )
            
            mermaid.to_png(output_file)
            resource['output'] = output_file
        
        def _render_vega(self, resource: Dict) -> None:
            """Render Vega/Vega-Lite charts to high-resolution PNG"""
            output_file = self.output_dir / f"vega_{resource['hash']}.png"
            spec = json.loads(resource['content'])
            if spec.get('$schema', '').startswith('https://vega.github.io/schema/vega-lite'):
                # For Vega-Lite, use vl_convert
                png_data = vlc.vegalite_to_png(vl_spec=spec, scale=1, ppi=300)
            else:
                png_data = vlc.vega_to_png(vg_spec=spec, scale=1, ppi=300)
            with open(output_file, "wb") as f:
                f.write(png_data)
            resource['output'] = output_file

    class ASTTransformer:
        """Transform Pandoc AST, replace diagrams and apply formatting"""
        
        def __init__(self, resources: List[Dict]):
            self.resources = {res['position']: res for res in resources}
        
        def transform(self, ast: Dict) -> Dict:
            """Transform AST by replacing diagrams and applying formatting"""
            if not isinstance(ast.get('blocks'), list):
                return ast
                
            for i, block in enumerate(ast['blocks']):
                if not isinstance(block, dict):
                    continue
                    
                pos = f"block_{i}"
                if pos in self.resources:
                    self._transform_diagram_block(ast['blocks'], i, pos)
                    
            return ast
        
        def _transform_diagram_block(self, blocks: List, index: int, position: str) -> None:
            """Replace diagram block with image or error node"""
            resource = self.resources[position]
            blocks[index] = (
                self._create_image_node(resource) if resource.get('output')
                else self._create_error_node(resource)
            )
        
        @staticmethod
        def _create_image_node(resource: Dict) -> Dict:
            """Create AST image node"""
            return {
                "t": "Para",
                "c": [{
                    "t": "Image",
                    "c": [
                        ["", [], []],
                        [{"t": "Str", 'c': resource.get('type', 'diagram')}],
                        [str(resource['output']), ""]
                    ]
                }]
            }
        
        @staticmethod
        def _create_error_node(resource: Dict) -> Dict:
            """Create AST error node for failed diagram rendering.
            
            Args:
                resource: Dictionary containing error details with keys:
                    - type: Diagram type (e.g., 'plantuml', 'mermaid')
                    - content: Original diagram content
                    
            Returns:
                AST node structure in Pandoc format
            """
            return {
                "t": "Div",
                "c": [
                    ["", ["error"], []],
                    [
                        {"t": "Para", "c": [
                            {"t": "Strong", "c": [{"t": "Str", "c": "Diagram Rendering Failed"}]},
                            {"t": "Space"},
                            {"t": "Str", "c": f"({resource.get('type', 'unknown')})"}
                        ]},
                        {"t": "CodeBlock", "c": [["", ["raw"], []], resource.get('content', '')]}
                    ]
                ]
            }

    class MarkdownConverter:
        """Main conversion processor"""

        def __init__(self, input_path: str, output_path: Optional[str] = None, logo_path: Optional[str] = 'static/logo.png'):
            self.input_path = Path(input_path)
            self.output_path = Path(output_path) if output_path else self.input_path.with_suffix('.docx')
            self.logo_path = logo_path
            self.working_dir = Path(tempfile.mkdtemp())
            self.resources = []
        
        
        def _find_diagrams(self, ast: Dict) -> None:
            """Locate diagrams in AST"""
            for i, block in enumerate(ast.get('blocks', [])):
                if not isinstance(block, dict) or block.get('t') != 'CodeBlock':
                    continue
                    
                lang = ''
                if len(block['c']) > 0 and isinstance(block['c'][0], list):
                    lang_info = block['c'][0]
                    if len(lang_info) > 1 and isinstance(lang_info[1], list) and len(lang_info[1]) > 0:
                        lang = lang_info[1][0]

                content = block['c'][1] if len(block['c']) > 1 else ''
                
                if lang in ['mermaid', 'vega', 'vega-lite']:
                    self.resources.append({
                        'type': lang,
                        'content': content,
                        'hash': hashlib.md5(content.encode()).hexdigest()[:8],
                        'position': f"block_{i}"
                    })
        
        def _generate_docx(self, ast: Dict) -> None:
            """Generate final Word document"""
            self.output_path.parent.mkdir(parents=True, exist_ok=True)
            
            args = ['--standalone', f'--resource-path={self.working_dir}']
            
            pypandoc.convert_text(
                json.dumps(ast),
                to='docx',
                format='json',
                outputfile=str(self.output_path),
                extra_args=args
            )

            document = Document(str(self.output_path)) 
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    # Detect if paragraph is a heading
                    is_heading = paragraph.style.name.startswith('Heading')
                    
                    # Apply different font settings based on paragraph type
                    if is_heading:
                        run.font.name = 'Arial'
                        run.font.bold = True
                        # Keep existing heading size
                    else:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
            self._apply_header_footer(document)
            document.save(str(self.output_path))
            
        # Function to copy headers and footers
        def _apply_header_footer(self, doc):
            """Apply header and footer from a reference document if provided"""
            header = doc.sections[0].header
            # Set section margins
            # Set header margins
            header_paragraph = header.paragraphs[0]
            header_paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.CENTER (0=left, 1=center, 2=right)

            logo_run = header_paragraph.add_run()
            logo_run.add_picture(self.logo_path, width=Inches(0.96))
            # Lower the image by adjusting its vertical position
            # Add footer
            footer = doc.sections[0].footer
            footer_paragraph = footer.paragraphs[0]  # Use the first header_paragraph in the footer
            footer_paragraph.alignment = 1  # Center alignment
            footer_paragraph.text = "NOAH AI-AI Agent Specialized in Life Science"
            if self.logo_path == 'static/roche-logo.png':
                footer_paragraph.text = "AI 智慧联  轻松做科研"
            footer_paragraph.style.font.size = Pt(9)  # Optional: adjust font size

        # Update the convert method to accept reference document
        def convert(self) -> bool:
            """Run the complete conversion process"""
            try:
                with open(self.input_path, 'r', encoding='utf-8') as f:
                    txt = f.read()
                lines = txt.splitlines()
                # Process text to ensure any line ending with ** and having matching ** has two newlines after it
                processed_lines = []

                for i, line in enumerate(lines):
                    processed_lines.append(line)
                    
                    # Check if line contains a complete ** pair and ends with **
                    if '**' in line and line.strip().endswith('**'):
                        # Find position of first ** in the line
                        first_pos = line.find('**')
                        # Find position of last ** in the line (which should be at the end)
                        last_pos = line.rfind('**')
                        
                        # Ensure there are actually two different ** markers
                        if first_pos < last_pos:
                            # Check if we need to add newlines (if not already two newlines after)
                            if i + 1 >= len(lines) or not lines[i+1].strip() == '':
                                processed_lines.append('')
                txt = '\n'.join(processed_lines)
                
                ast = json.loads(pypandoc.convert_text(
                    txt,
                    to='json',
                    format='markdown'
                ))
                
                self._find_diagrams(ast)
                MarkdownToWordTool.DiagramRenderer(self.working_dir).render(self.resources)
                ast = MarkdownToWordTool.ASTTransformer(self.resources).transform(ast)
                
                self._generate_docx(ast)
                
                    
                return True
            except Exception as e:
                print(f"Conversion failed: {str(e)}")
                return False
            finally:
                shutil.rmtree(self.working_dir, ignore_errors=True)
