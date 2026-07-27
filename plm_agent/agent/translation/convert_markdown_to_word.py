"""
Markdown -> Word converter used by translation pipeline.
Migrated from archived `code/markdown_to_word` for production use.

字体说明：使用的 宋体、黑体、Times New Roman 均为系统/Office 常见字体，
无需在项目中附带字体文件；若运行环境无对应字体，Word 会使用系统默认替代。
"""

from __future__ import annotations

import html
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Optional
from urllib.request import Request, urlopen

from docx import Document

logger = logging.getLogger(__name__)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Fallback maximum image width when page dimensions are unavailable (A4 content area)
_DEFAULT_MAX_IMAGE_WIDTH_INCHES = 6.0


class MarkdownToWordConverter:
    def __init__(
        self,
        base_dir: Optional[str] = None,
        format_type: str = "english",
        image_url_to_path: Optional[Dict[str, str]] = None,
    ):
        self.doc = Document()
        self.base_dir = Path(base_dir) if base_dir else None
        self.format_type = (format_type or "english").lower()
        self.image_url_to_path = image_url_to_path or {}
        self._setup_styles()

    def _setup_styles(self) -> None:
        if self.format_type == "chinese":
            self._setup_chinese_styles()
        else:
            self._setup_english_styles()

    def _setup_english_styles(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._create_english_heading_styles()

    def _setup_chinese_styles(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._create_chinese_heading_styles()

    def _create_english_heading_styles(self) -> None:
        """预定义英文标题与 Caption 样式（与 code/markdown_to_word 一致）。"""
        black = RGBColor(0, 0, 0)
        for level, size_pt, bold, italic in [
            (1, 14, True, False),
            (2, 12, True, False),
            (3, 12, False, True),
        ]:
            try:
                s = self.doc.styles[f"Heading {level}"]
            except KeyError:
                s = self.doc.styles["Normal"]
            f = s.font
            f.name = "Times New Roman"
            f._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            f.size = Pt(size_pt)
            f.bold = bold
            f.italic = italic
            f.color.rgb = black
            if hasattr(s, "paragraph_format"):
                s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            cap = self.doc.styles["Caption"]
        except KeyError:
            cap = self.doc.styles["Normal"]
        f = cap.font
        f.name = "Times New Roman"
        f._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        f.size = Pt(10)
        if hasattr(cap, "paragraph_format"):
            cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_chinese_heading_styles(self) -> None:
        """预定义中文标题与 Caption 样式（黑体/宋体、段前段后 6pt，与 code/markdown_to_word 一致）。"""
        black = RGBColor(0, 0, 0)
        space_pt = Pt(6)
        # H1: 黑体 15pt 加粗；H2: 黑体 14pt 加粗；H3: 黑体 12pt 加粗；H4: 宋体 12pt 加粗
        for level, font_name, size_pt in [
            (1, "黑体", 15),
            (2, "黑体", 14),
            (3, "黑体", 12),
            (4, "宋体", 12),
            (5, "宋体", 12),
            (6, "宋体", 12),
        ]:
            try:
                s = self.doc.styles[f"Heading {level}"]
            except KeyError:
                s = self.doc.styles["Normal"]
            f = s.font
            f.name = font_name
            f._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            f.size = Pt(size_pt)
            f.bold = True
            f.color.rgb = black
            if hasattr(s, "paragraph_format"):
                s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                s.paragraph_format.space_before = space_pt
                s.paragraph_format.space_after = space_pt
        try:
            cap = self.doc.styles["Caption"]
        except KeyError:
            cap = self.doc.styles["Normal"]
        f = cap.font
        f.name = "宋体"
        f._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        f.size = Pt(10.5)
        if hasattr(cap, "paragraph_format"):
            cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _strip_markdown_symbols(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"\s*#+\s*$", "", text)
        text = re.sub(r"^\s*\*+\s*", "", text)
        text = re.sub(r"\s*\*+\s*$", "", text)
        return text.strip()

    @staticmethod
    def _strip_cell_markdown(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"(?<!_)_(?!_)([^_]+?)(?<!_)_(?!_)", r"\1", text)
        return html.unescape(text).strip()

    def _parse_markdown_lines(self, content: str):
        lines = content.split("\n")
        parsed = []
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                parsed.append(("heading", min(len(heading_match.group(1)), 6), heading_match.group(2)))
                i += 1
                continue

            # 整行 HTML 标题：<h1>..</h1> 到 <h6>..</h6>
            html_heading_match = re.match(r"^\s*<h([1-6])>\s*(.*?)\s*</h\1>\s*$", line, re.IGNORECASE | re.DOTALL)
            if html_heading_match:
                level = min(max(int(html_heading_match.group(1)), 1), 6)
                inner = html_heading_match.group(2).strip()
                parsed.append(("heading", level, inner))
                i += 1
                continue

            bold_numbered_match = re.match(r"^\*\*(\d+(?:\.\d+)*)\s+(.+?)\*\*\s*$", line)
            if bold_numbered_match and "." in bold_numbered_match.group(1):
                nums = bold_numbered_match.group(1).split(".")
                level = 2 if len(nums) == 2 else 3
                parsed.append(("heading", level, f"{bold_numbered_match.group(1)} {bold_numbered_match.group(2)}"))
                i += 1
                continue

            numbered_heading_match = re.match(r"^(\d+(?:\.\d+)+)\s+(.+)$", line)
            if numbered_heading_match and not line.strip().startswith("**"):
                nums = numbered_heading_match.group(1).split(".")
                level = 2 if len(nums) == 2 else 3
                parsed.append(("heading", level, f"{numbered_heading_match.group(1)} {numbered_heading_match.group(2)}"))
                i += 1
                continue

            if "|" in line and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r"^[\|\s\-:+]+$", next_line):
                    table_lines = [line, lines[i + 1]]
                    i += 2
                    while i < len(lines) and "|" in lines[i] and not re.match(r"^[\|\s\-:+]+$", lines[i].strip()):
                        table_lines.append(lines[i])
                        i += 1
                    parsed.append(("table", table_lines, None))
                    continue

            img_match = re.match(r"!\[([^\]]*)\]\(([^\)]+)\)", line)
            if img_match:
                parsed.append(("image", img_match.group(1), img_match.group(2), None))
                i += 1
                continue

            caption_match = re.match(r"^.*?(?:Figure|Table|图|表|FIGURE|TABLE)\s*\d+[\.:]?\s*(.+)$", line, re.IGNORECASE)
            if caption_match and parsed:
                prev = parsed[-1]
                if prev[0] == "image":
                    parsed[-1] = ("image", prev[1], prev[2], line.strip())
                    i += 1
                    continue
                if prev[0] == "table":
                    parsed[-1] = ("table", prev[1], line.strip())
                    i += 1
                    continue

            if line.startswith("```"):
                code_lines = [line[3:].strip()]
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1
                parsed.append(("code", code_lines))
                continue

            if line.strip():
                parsed.append(("paragraph", line))
            else:
                parsed.append(("blank",))
            i += 1
        return parsed

    def _parse_table(self, table_lines):
        if len(table_lines) < 2:
            return None
        sep_pattern = re.compile(r"^[\|\s\-:+]+$")
        rows = []
        for line in table_lines:
            s = line.strip()
            if sep_pattern.match(s):
                continue
            cells = [c.strip() for c in line.split("|")]
            if cells and cells[0] == "" and cells[-1] == "":
                cells = cells[1:-1]
            rows.append([self._strip_cell_markdown(c) for c in cells])
        if not rows:
            return None
        max_cols = max(len(r) for r in rows)
        return [r + [""] * (max_cols - len(r)) if len(r) < max_cols else r[:max_cols] for r in rows]

    def _apply_run_font(self, run, size_pt: float = 12, color: Optional[RGBColor] = None):
        if self.format_type == "chinese":
            run.font.name = "宋体"
            run.font._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        else:
            run.font.name = "Times New Roman"
            run.font._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size_pt)
        if color is not None:
            run.font.color.rgb = color

    def _add_heading(self, level: int, text: str) -> None:
        text = html.unescape(self._strip_markdown_symbols(text))
        style = f"Heading {min(max(level, 1), 6)}"
        p = self.doc.add_paragraph(text, style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        black = RGBColor(0, 0, 0)
        if self.format_type == "chinese":
            for run in p.runs:
                run.font.color.rgb = black
        else:
            for run in p.runs:
                run.font.color.rgb = black
                if level == 1:
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.italic = False
                elif level == 2:
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.italic = False
                elif level == 3:
                    run.font.size = Pt(12)
                    run.font.bold = False
                    run.font.italic = True
                elif level >= 4:
                    run.font.size = Pt(12)
                    run.font.bold = False
                    run.font.italic = False
                run.font.name = "Times New Roman"
                run.font._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def _add_paragraph(self, text: str) -> None:
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.format_type == "chinese" else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            self._apply_run_font(run)

    def _add_table(self, rows, caption: Optional[str] = None) -> None:
        if not rows:
            return
        if caption:
            cp = self.doc.add_paragraph(html.unescape(caption), style="Caption")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                self._apply_run_font(run, 10.5 if self.format_type == "chinese" else 10, RGBColor(0, 0, 0))
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._apply_run_font(run, 12)

    def _content_width_inches(self) -> float:
        """Return the usable content width in inches from the document's first section."""
        try:
            sec = self.doc.sections[0]
            emu = sec.page_width - sec.left_margin - sec.right_margin
            return max(emu / 914400, 1.0)
        except Exception:
            return _DEFAULT_MAX_IMAGE_WIDTH_INCHES

    def _add_image(self, alt_text: str, img_path: str, caption: Optional[str] = None) -> None:
        alt_text = html.unescape(alt_text or "")
        caption = html.unescape(caption or "")
        img_path = (img_path or "").strip()

        # 若有 url→本地路径 映射则直接读本地插入，不下载
        if img_path and self.image_url_to_path and img_path in self.image_url_to_path:
            local_path = Path(self.image_url_to_path[img_path])
            if local_path.is_file():
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(local_path), width=Inches(self._content_width_inches()))
                if caption or alt_text:
                    cp = self.doc.add_paragraph(caption or alt_text, style="Caption")
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cp.runs:
                        self._apply_run_font(
                            run, 10.5 if self.format_type == "chinese" else 10, RGBColor(0, 0, 0)
                        )
                return

        # 再尝试按 URL 下载并插入（MD 中常见 ![alt](https://...)）
        if img_path.lower().startswith(("http://", "https://")):
            try:
                req = Request(img_path, headers={"User-Agent": "Mozilla/5.0 (compatible; NoahAgent/1.0)"})
                with urlopen(req, timeout=30) as resp:
                    data = resp.read()
                if data:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(BytesIO(data), width=Inches(self._content_width_inches()))
                    if caption or alt_text:
                        cp = self.doc.add_paragraph(caption or alt_text, style="Caption")
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cp.runs:
                            self._apply_run_font(
                                run, 10.5 if self.format_type == "chinese" else 10, RGBColor(0, 0, 0)
                            )
                    return
            except Exception as e:
                logger.warning("[md2docx] download image failed url=%s err=%s", img_path[:80], e)

        # 再按本地路径查找
        img = Path(img_path)
        abs_path = None
        if img.is_absolute() and img.exists():
            abs_path = img
        elif img.exists():
            abs_path = img.resolve()
        elif self.base_dir and (self.base_dir / img_path).exists():
            abs_path = (self.base_dir / img_path).resolve()
        if abs_path is None:
            self._add_paragraph(f"[图片: {alt_text if alt_text else img_path}]")
            if caption:
                cp = self.doc.add_paragraph(caption, style="Caption")
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(abs_path), width=Inches(self._content_width_inches()))
        if caption or alt_text:
            cp = self.doc.add_paragraph(caption or alt_text, style="Caption")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                self._apply_run_font(run, 10.5 if self.format_type == "chinese" else 10, RGBColor(0, 0, 0))

    def _add_code_block(self, code_lines) -> None:
        code_text = "\n".join(code_lines[1:]) if len(code_lines) > 1 else "\n".join(code_lines)
        code_text = html.unescape(code_text)
        p = self.doc.add_paragraph(code_text)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    def _add_page_footer(self) -> None:
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        fld2 = OxmlElement("w:instrText")
        fld2.set(qn("xml:space"), "preserve")
        fld2.text = "PAGE"
        fld3 = OxmlElement("w:fldChar")
        fld3.set(qn("w:fldCharType"), "end")
        run._element.append(fld1)
        run._element.append(fld2)
        run._element.append(fld3)
        self._apply_run_font(run, 10.5 if self.format_type == "chinese" else 10, RGBColor(0, 0, 0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _process_paragraph_with_formatting(self, text: str) -> None:
        text = html.unescape(text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.format_type == "chinese" else WD_ALIGN_PARAGRAPH.LEFT

        def apply_formatting(
            s: str,
            *,
            bold: bool = False,
            italic: bool = False,
            superscript: bool = False,
            subscript: bool = False,
        ) -> None:
            if not s:
                return
            # 内联 HTML 标签：<sup>, <sub>, <b>, <strong>, <i>, <em>（优先于 Markdown 语法）
            candidates = []
            for pattern, kind in [
                (r"<sup>([^<]*)</sup>", "sup"),
                (r"<sub>([^<]*)</sub>", "sub"),
                (r"<b>(.*?)</b>", "b"),
                (r"<strong>(.*?)</strong>", "strong"),
                (r"<i>(.*?)</i>", "i"),
                (r"<em>(.*?)</em>", "em"),
            ]:
                m = re.search(pattern, s, re.IGNORECASE | (re.DOTALL if kind in ("b", "strong", "i", "em") else 0))
                if m:
                    candidates.append((m.start(), m, kind))
            if candidates:
                _, first_m, kind = min(candidates, key=lambda x: x[0])
                apply_formatting(s[: first_m.start()], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                if kind == "sup":
                    run = p.add_run(first_m.group(1))
                    run.font.superscript = True
                    self._apply_run_font(run)
                elif kind == "sub":
                    run = p.add_run(first_m.group(1))
                    run.font.subscript = True
                    self._apply_run_font(run)
                elif kind in ("b", "strong"):
                    apply_formatting(first_m.group(1), bold=True, italic=italic, superscript=superscript, subscript=subscript)
                else:
                    apply_formatting(first_m.group(1), bold=bold, italic=True, superscript=superscript, subscript=subscript)
                apply_formatting(s[first_m.end() :], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                return

            link_match = re.search(r"\[([^\]]+)\]\(([^\)]+)\)", s)
            if link_match:
                apply_formatting(s[: link_match.start()], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                link_text, url = link_match.group(1), link_match.group(2)
                run = p.add_run(link_text)
                run.bold = bold
                run.italic = italic
                run.font.superscript = superscript
                run.font.subscript = subscript
                self._apply_run_font(run, 12, RGBColor(0, 0, 255))
                try:
                    r_id = p.part.relate_to(
                        url,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                        is_external=True,
                    )
                    hyperlink = OxmlElement("w:hyperlink")
                    hyperlink.set(qn("r:id"), r_id)
                    run._element.getparent().remove(run._element)
                    hyperlink.append(run._element)
                    p._element.append(hyperlink)
                except Exception:
                    pass
                apply_formatting(s[link_match.end() :], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                return

            bold_match = re.search(r"\*\*([^*]+)\*\*|__(?!_)([^_]+)__(?!_)", s)
            if bold_match:
                apply_formatting(s[: bold_match.start()], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                val = bold_match.group(1) if bold_match.group(1) else bold_match.group(2)
                run = p.add_run(val)
                run.bold = True
                run.italic = italic
                run.font.superscript = superscript
                run.font.subscript = subscript
                self._apply_run_font(run)
                apply_formatting(s[bold_match.end() :], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                return

            italic_match = re.search(r"(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)([^_]+?)(?<!_)_(?!_)", s)
            if italic_match:
                apply_formatting(s[: italic_match.start()], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                val = italic_match.group(1) if italic_match.group(1) else italic_match.group(2)
                run = p.add_run(val)
                run.bold = bold
                run.italic = True
                run.font.superscript = superscript
                run.font.subscript = subscript
                self._apply_run_font(run)
                apply_formatting(s[italic_match.end() :], bold=bold, italic=italic, superscript=superscript, subscript=subscript)
                return

            run = p.add_run(s)
            run.bold = bold
            run.italic = italic
            run.font.superscript = superscript
            run.font.subscript = subscript
            self._apply_run_font(run)

        apply_formatting(text)

    def convert(self, markdown_content: str, output_path: str) -> None:
        for item in self._parse_markdown_lines(markdown_content):
            kind = item[0]
            if kind == "heading":
                self._add_heading(item[1], item[2])
            elif kind == "paragraph":
                self._process_paragraph_with_formatting(item[1])
            elif kind == "blank":
                self.doc.add_paragraph()
            elif kind == "table":
                rows = self._parse_table(item[1])
                self._add_table(rows, item[2] if len(item) > 2 else None)
            elif kind == "image":
                self._add_image(item[1], item[2], item[3] if len(item) > 3 else None)
            elif kind == "code":
                self._add_code_block(item[1])

        self._add_page_footer()
        self.doc.save(output_path)


def convert_markdown_to_word(
    input_file: str,
    output_file: Optional[str] = None,
    format_type: str = "english",
    *,
    image_url_to_path: Optional[Dict[str, str]] = None,
) -> str:
    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"找不到输入文件: {input_file}")

    markdown_content = input_path.read_text(encoding="utf-8")
    out = Path(output_file) if output_file else input_path.with_suffix(".docx")
    converter = MarkdownToWordConverter(
        base_dir=str(input_path.parent),
        format_type=format_type,
        image_url_to_path=image_url_to_path,
    )
    converter.convert(markdown_content, str(out))
    return str(out)

