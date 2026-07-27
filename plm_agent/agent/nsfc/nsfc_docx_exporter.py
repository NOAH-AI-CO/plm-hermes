from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Optional
import re
import os
import logging

logger = logging.getLogger(__name__)


class NSFCDocxExporter:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self.heading_style_map = {} 
        self.template_path = template_path 
        self.original_template_texts = set()
        
        for para in self.doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                self.original_template_texts.add(text)
        
        self._init_normal_style()
    
    @staticmethod
    def _count_matching_titles(template_path: str, markdown_content: str) -> int:
        try:
            temp_exporter = NSFCDocxExporter(template_path)
            parts = temp_exporter._parse_markdown_to_parts(markdown_content)

            matched = 0
            for part in parts:
                title = part.get('title', '').strip()
                if title:
                    idx = temp_exporter._find_paragraph_index_by_keyword(title)
                    if idx is not None:
                        matched += 1
            return matched
        except Exception as e:
            logger.warning(f"测试模板 {template_path} 时出错: {e}")
            return 0
    
    @staticmethod
    def auto_select_template(markdown_path: str) -> str:
        with open(markdown_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 获取模板目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        young_template = os.path.join(current_dir, 'young_template.docx')
        general_template = os.path.join(current_dir, 'general_template.docx')
        
        young_matched = NSFCDocxExporter._count_matching_titles(young_template, md_content)
        logger.info(f"   青年基金模板匹配: {young_matched} 个标题")
        
        general_matched = NSFCDocxExporter._count_matching_titles(general_template, md_content)
        logger.info(f"   面上项目模板匹配: {general_matched} 个标题")
        
        # 选择匹配度更高的模板
        if young_matched >= general_matched:
            logger.info(f"自动选择: 青年基金模板 (匹配度: {young_matched} vs {general_matched})")
            return young_template
        else:
            logger.info(f"自动选择: 面上项目模板 (匹配度: {general_matched} vs {young_matched})")
            return general_template

    # ================== 基础样式 ==================
    def _init_normal_style(self):
        """
        初始化样式
        注意：不修改 Normal 样式，避免影响模板原有段落
        只初始化标题样式，正文格式在插入时单独设置
        """
        # 只初始化标题样式，不修改 Normal 样式
        # 正文格式会在 _write_content_into_paragraph 中单独设置
        self._init_heading_styles()
    
    def _init_heading_styles(self):
        """
        NSFC 标题格式要求：
        - 宋体、小四号（12pt）
        - 加粗
        - 无首行缩进
        - 段前段后6磅间距
        """
        styles = self.doc.styles
        
        self.heading_style_map = {}
        
        for i in range(2, 6):
            try:
                style_name = f"NSFCHeading{i}"
                heading = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                logger.info(f"创建新标题样式: {style_name}")
            except:
                logger.warning(f"无法创建标题样式 {style_name}")
                continue
            
            if heading and style_name:
                # 保存样式映射
                self.heading_style_map[i] = style_name
                
                # 字体设置
                heading.font.size = Pt(12)
                heading.font.name = "Times New Roman"
                heading.font.bold = True
                
                if hasattr(heading, '_element') and hasattr(heading._element, 'rPr'):
                    rFonts = heading._element.rPr.rFonts
                    rFonts.set(qn("w:ascii"), "Times New Roman")
                    rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    rFonts.set(qn("w:eastAsia"), "宋体")
                    rFonts.set(qn("w:cs"), "Times New Roman")
                
                # 段落格式
                pfmt = heading.paragraph_format
                pfmt.first_line_indent = Pt(0)  # 无首行缩进
                pfmt.left_indent = Pt(0)
                pfmt.space_before = Pt(6)  # 段前6磅
                pfmt.space_after = Pt(6)   # 段后6磅
                pfmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pfmt.line_spacing = Pt(22)
                pfmt.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        logger.info(f"标题样式初始化完成，映射: {self.heading_style_map}")


    def _normalize_title(self, title: str) -> str:
        if not title:
            return ""

        s = str(title)
        s = re.sub(r"\s+", "", s)
        s = re.sub(
            r"[（(][^）)]*[\u4e00-\u9fa5]{2,}[^（(]*[）)]",
            "",
            s,
        )
        # 删除开头的编号（支持半角点号.、全角点号．、顿号、、括号等）
        s = re.sub(
            r"^[（(]?[一二三四五六七八九十0-9]+[)）\.．、]",
            "",
            s,
        )
        s = s.rstrip("。；：、.;；")
        
        return s

    def _find_paragraph_index_by_keyword(self, keyword: str) -> Optional[int]:
        keyword = (keyword or "").strip()
        if not keyword:
            return None

        normalized_keyword = self._normalize_title(keyword)

        # 1. 精确匹配
        for i, p in enumerate(self.doc.paragraphs):
            text = (p.text or "").strip()
            if not text:
                continue
            if keyword in text:
                logger.debug(f" 精确匹配标题: '{keyword}' 在段落 {i}")
                return i

        # 2. 标准化匹配
        for i, p in enumerate(self.doc.paragraphs):
            text = (p.text or "").strip()
            if not text:
                continue
            normalized_text = self._normalize_title(text)
            if normalized_keyword and normalized_keyword in normalized_text:
                logger.info(
                    f" 标准化匹配成功: '{keyword}' ≈ '{text}' (段落 {i})"
                )
                return i

        # 3. 模糊匹配（对于长标题，匹配前20个字符，避免细微差异）
        if len(keyword) > 30:
            keyword_prefix = keyword[:20]
            for i, p in enumerate(self.doc.paragraphs):
                text = (p.text or "").strip()
                if not text or len(text) < 15:
                    continue
                # 匹配前20个字符
                if text.startswith(keyword_prefix):
                    logger.info(
                        f" 模糊匹配成功(前20字符): '{keyword[:30]}...' ≈ '{text[:30]}...' (段落 {i})"
                    )
                    return i

        logger.warning(f"未找到标题: '{keyword[:80]}...' (标准化后: '{normalized_keyword[:80]}...')")
        return None


    def _is_subheading_line(self, text: str) -> bool:
        if not text:
            return False
        t = text.strip()

        if len(t) >= 3 and t[0] in {"（", "("} and t[2] in {"）", ")"} and t[1].isdigit():
            return True
        if len(t) >= 2 and t[0].isdigit() and t[1] in {".", "．", "、"}:
            return True
        if re.match(r"^[一二三四五六七八九十]+、", t):
            return True

        return False

    def _find_body_paragraph_after_title(self, title_index: int) -> Optional[int]:
        """
        找到标题后第一个应该写入内容的段落位置。
        注意：蓝色说明文字（如"为什么要开展此项研究"）应该跳过，不能修改其格式。
        """
        paras = self.doc.paragraphs
        n = len(paras)

        for i in range(title_index + 1, n):
            p = paras[i]
            text = (p.text or "").strip()
            if not text:
                return i
            if self._is_subheading_line(text):
                continue
            if text in self.original_template_texts:
                continue
            return i
        return None


    def _is_heading_line(self, text: str) -> bool:
        """判断是否为标题行（多级数字标题，如 1.1、2.3.4）"""
        if not text:
            return False
        t = text.strip()
        
        # 只匹配多级数字标题：1.1、1.2.3、2.3.4.5 等
        # 必须满足以下条件：
        # 1. 以多级数字开头（如 1.1、2.3.4）
        # 2. 文本较短（< 80字符，纯标题行）
        # 3. 不包含句号（排除正文）
        if re.match(r"^(\d+\.){1,}\d+\s+", t):
            # 只有短文本才是纯标题行（排除"标题+正文"的混合段落）
            if len(t) < 80 and "。" not in t:
                return True
        
        return False
    
    def _get_heading_level(self, text: str) -> int:
        """
        判断标题级别（返回 0 表示非标题）
        
        返回值：
        - 0: 非标题（普通段落）
        - 1: 一级标题（如 1.、2.、3.）
        - 2: 二级标题（如 1.1、1.2、2.1）
        - 3: 三级标题（如 1.1.1、1.2.1、2.1.1）
        - 4+: 更深层级标题
        """
        if not text:
            return 0
        
        t = text.strip()
        
        # 匹配多级数字标题：1.、1.1、1.1.1 等
        match = re.match(r"^((\d+\.)+)\d+\s+", t)
        if match:
            # 只有短文本才是纯标题行（排除"标题+正文"的混合段落）
            if len(t) < 80 and "。" not in t:
                # 计算点号的数量来确定级别
                dots = match.group(1)
                level = dots.count('.')
                return level
        
        return 0
    
    def _should_break_paragraph(self, text: str) -> bool:
        """判断是否应该断段（让某些行单独成段，以保证首行缩进）"""
        if not text:
            return False
        t = text.strip()
        
        # 匹配"目标X："、"（X）"等模式，这些应该单独成段以保证每行都有首行缩进
        # 如：目标一：、目标二：、（1）、（2）等
        if re.match(r"^目标[一二三四五六七八九十]+：", t):
            return True
        if re.match(r"^[（(]\d+[）)]", t):
            return True
        
        return False
    
    def _is_bold_heading(self, text: str) -> bool:
        """判断是否是 Markdown 加粗标题（如 **1. 研究意义**）"""
        if not text:
            return False
        t = text.strip()
        
        # 匹配 **数字. 标题** 格式
        if re.match(r"^\*\*\s*\d+\.\s+.+\*\*$", t):
            return True
        # 匹配 **数字、标题** 格式
        if re.match(r"^\*\*\s*\d+[、．]\s+.+\*\*$", t):
            return True
        
        return False
    
    def _remove_bold_markers(self, text: str) -> str:
        """移除 Markdown 加粗标记 ** """
        if not text:
            return text
        # 移除前后的 **
        return text.strip().strip('*').strip()
    
    def _clean_markdown_escapes(self, text: str) -> str:
        """
        清理Markdown转义字符，主要处理参考文献引用
        例如：\[1] -> [1], \[2\] -> [2]
        """
        if not text:
            return text
        
        # 替换转义的方括号
        text = text.replace(r"\[", "[")
        text = text.replace(r"\]", "]")
        
        return text

    def _write_content_into_paragraph(self, paragraph_index: int, content: str):
        if content is None:
            return
        content = content.strip()
        if not content:
            return
        
        # 清理Markdown转义字符（如参考文献的 \[1] -> [1]）
        content = self._clean_markdown_escapes(content)

        anchor_para = self.doc.paragraphs[paragraph_index]

        p = anchor_para._element
        for r in list(anchor_para.runs):
            p.remove(r._element)

        paragraphs = []
        buf = []

        for raw_line in content.split("\n"):
            line = raw_line.rstrip()
            stripped = line.strip()
            
            if stripped == "":
                if buf:
                    paragraphs.append("\n".join(buf).strip())
                    buf = []
            else:
                if (self._is_heading_line(stripped) or self._should_break_paragraph(stripped)) and buf:
                    paragraphs.append("\n".join(buf).strip())
                    buf = []
                    paragraphs.append(stripped)  # 当前行单独成段
                elif self._is_heading_line(stripped) or self._should_break_paragraph(stripped):
                    paragraphs.append(stripped)
                # 普通行加入缓冲区
                else:
                    buf.append(line)

        if buf:
            paragraphs.append("\n".join(buf).strip())

        if not paragraphs:
            return

        for para_text in paragraphs:
            p_new = anchor_para.insert_paragraph_before()
            
            # 检查是否是 Markdown 加粗标题（如 **1. 研究意义**）
            is_bold_heading = self._is_bold_heading(para_text)
            
            # 判断数字标题级别（如 1.1、1.2.3）
            heading_level = self._get_heading_level(para_text)
            
            # 处理 Markdown 加粗标题
            if is_bold_heading:
                # 移除 ** 标记
                clean_text = self._remove_bold_markers(para_text)
                
                # 设置为无首行缩进的 Normal 样式
                p_new.style = self.doc.styles["Normal"]
                p_new.paragraph_format.first_line_indent = Pt(0)
                p_new.paragraph_format.left_indent = Pt(0)
                p_new.paragraph_format.space_before = Pt(6)
                p_new.paragraph_format.space_after = Pt(6)
                
                # 添加文本并设置加粗
                run = p_new.add_run(clean_text)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
                run.font.size = Pt(12)
                run.font.bold = True
                
            # 处理数字标题（如 1.1、1.2.3）
            elif heading_level > 0 and hasattr(self, 'heading_style_map'):
                style_level = heading_level + 1  # level 1 -> Heading 2
                style_name = self.heading_style_map.get(style_level)
                
                if style_name:
                    try:
                        p_new.style = self.doc.styles[style_name]
                        logger.debug(f"使用标题样式 {style_name} for {para_text[:30]}...")
                    except KeyError:
                        # 如果样式不存在，降级使用 Normal 样式并手动设置格式
                        logger.debug(f"标题样式 {style_name} 不可用，手动设置格式")
                        p_new.style = self.doc.styles["Normal"]
                        p_new.paragraph_format.first_line_indent = Pt(0)
                        p_new.paragraph_format.left_indent = Pt(0)
                        p_new.paragraph_format.space_before = Pt(6)
                        p_new.paragraph_format.space_after = Pt(6)
                else:
                    # 没有找到样式映射，手动设置格式
                    p_new.style = self.doc.styles["Normal"]
                    p_new.paragraph_format.first_line_indent = Pt(0)
                    p_new.paragraph_format.left_indent = Pt(0)
                    p_new.paragraph_format.space_before = Pt(6)
                    p_new.paragraph_format.space_after = Pt(6)
                
                run = p_new.add_run(para_text)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
                run.font.size = Pt(12)
                run.font.bold = True
                
            # 普通正文段落
            else:
                p_new.style = self.doc.styles["Normal"]
                
                # 设置正文格式（不影响模板原有段落）
                p_new.paragraph_format.first_line_indent = Pt(24)  # 首行缩进两字符
                p_new.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_new.paragraph_format.line_spacing = Pt(22)
                p_new.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                run = p_new.add_run(para_text)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
                run.font.size = Pt(12)

        parent = anchor_para._element.getparent()
        parent.remove(anchor_para._element)


    def _merge_subsections(self, parts: List[Dict[str, str]]) -> List[Dict[str, str]]:
        main_markers = ["（一）", "（二）", "（三）", "（四）", "（五）"]
        main_section_indices = {}
        for i, part in enumerate(parts):
            title = part.get("title", "").strip()
            for marker in main_markers:
                if marker in title:
                    main_section_indices[marker] = i
                    break
        
        merged = []
        current_main_section = None
        current_main_section_marker = None
        orphan_subsections = []  # 孤立的子章节（主章节还未出现）
        
        for i, part in enumerate(parts):
            title = part.get("title", "").strip()
            content = part.get("content", "").strip()
            
            is_main_section = False
            section_marker = None
            
            for marker in main_markers:
                if marker in title:
                    is_main_section = True
                    section_marker = marker
                    break
            
            is_numbered_item = bool(re.match(r'^[0-9]+[\.．、]', title))
            
            if is_main_section:
                if current_main_section:
                    merged.append(current_main_section)
                
                filtered_content_lines = []
                for line in content.split('\n'):
                    line_stripped = line.strip()
                    if line_stripped and line_stripped not in self.original_template_texts:
                        filtered_content_lines.append(line)
                filtered_content = '\n'.join(filtered_content_lines).strip()
                
                current_main_section = {
                    "title": title,
                    "content": filtered_content,  # 使用过滤后的content
                    "level": part.get("level", 2)
                }
                current_main_section_marker = section_marker
                
                if orphan_subsections and section_marker not in ["（三）", "（四）"]:
                    for orphan in orphan_subsections:
                        orphan_title = orphan.get("title", "").strip()
                        orphan_content = orphan.get("content", "").strip()
                        orphan_is_numbered = bool(re.match(r'^[0-9]+[\.．、]', orphan_title))
                        
                        if not orphan_is_numbered:
                            if current_main_section["content"]:
                                current_main_section["content"] += "\n\n"
                            if orphan_title:
                                current_main_section["content"] += f"**{orphan_title}**\n\n"
                            if orphan_content:
                                current_main_section["content"] += orphan_content
                    orphan_subsections = []
                
            else:
                if current_main_section:
                    if current_main_section_marker in ["（三）", "（四）"] and is_numbered_item:
                        merged.append(part)
                    else:
                        if current_main_section["content"]:
                            current_main_section["content"] += "\n\n"
                        if title:
                            current_main_section["content"] += f"**{title}**\n\n"
                        if content:
                            current_main_section["content"] += content
                else:
                    orphan_subsections.append(part)
        
        if current_main_section:
            merged.append(current_main_section)
        
        merged.extend(orphan_subsections)
        
        logger.info(f"合并子章节：原始 {len(parts)} 个 parts -> 合并后 {len(merged)} 个 parts")
        return merged

    def _fill_by_parts(self, parts: List[Dict[str, str]]):
        logger.info(f"开始填充 {len(parts)} 个章节到 Word 模板")

        merged_parts = self._merge_subsections(parts)
        success, failed = 0, []

        for part in merged_parts:
            raw_title = part.get("title") or ""
            content = part.get("content") or ""
            title_key = raw_title.strip()
            if not title_key:
                continue

            title_idx = self._find_paragraph_index_by_keyword(title_key)
            if title_idx is None:
                logger.warning(f" 模板中未找到标题：{title_key}")
                failed.append(title_key)
                continue

            body_idx = self._find_body_paragraph_after_title(title_idx)
            if body_idx is None:
                logger.warning(f" 未找到标题 [{title_key}] 对应的正文段落")
                failed.append(title_key)
                continue

            content_lines = content.split('\n')
            filtered_lines = []
            
            for line in content_lines:
                line_stripped = line.strip()
                if line_stripped and line_stripped not in self.original_template_texts:
                    filtered_lines.append(line)
                elif not line_stripped:  # 保留空行
                    filtered_lines.append(line)
            
            content = '\n'.join(filtered_lines).strip()
            
            if not content:
                logger.debug(f" 内容为空（说明文字已过滤）: {title_key}")
                success += 1
                continue

            self._write_content_into_paragraph(body_idx, content)
            success += 1
            logger.debug(f" 成功填充: {title_key[:50]}...")

        logger.info(f" Word 填充完成: 成功 {success}/{len(merged_parts)} 个章节")
        if failed:
            logger.warning(" 以下标题未能填充（最多显示 5 个）：")
            for t in failed[:5]:
                logger.warning("   - " + t)

    def _parse_markdown_to_parts(self, md_content: str) -> List[Dict[str, str]]:
        lines = md_content.split("\n")
        parts: List[Dict[str, str]] = []
        current_title: Optional[str] = None
        current_lines: List[str] = []

        def flush_current():
            nonlocal current_title, current_lines
            if current_title:
                content = "\n".join(current_lines).strip()
                parts.append(
                    {
                        "title": current_title,
                        "content": content,
                        "level": 2,
                    }
                )
            current_title = None
            current_lines = []

        for raw_line in lines:
            line = raw_line.rstrip("\n")
            stripped = line.strip()

            if stripped == "":
                if current_title is not None:
                    current_lines.append("")
                continue

            if stripped.startswith("---") or (
                stripped.startswith("**") and "：**" in stripped
            ):
                continue

            if stripped.startswith("#"):
                hash_count = 0
                for ch in stripped:
                    if ch == "#":
                        hash_count += 1
                    else:
                        break

                title_text = stripped[hash_count:].strip()
                if hash_count == 1:
                    flush_current()
                    continue
                if hash_count == 2 or hash_count == 3:
                    flush_current()
                    current_title = title_text
                    current_lines = []
                    continue
                if hash_count >= 4:
                    if current_title is not None:
                        current_lines.append(title_text)
                    continue
            if current_title is not None:
                current_lines.append(stripped)

        flush_current()
        logger.info(f"Markdown 解析完成，共 {len(parts)} 个顶层章节")
        for i, part in enumerate(parts):
            logger.debug(
                f"  [{i+1}] {part['title']} (len={len(part.get('content',''))})"
            )
        return parts

    def _fix_chapter_title_indent(self):
        import re
        for para in self.doc.paragraphs:
            original_text = para.text or ""
            text = original_text.strip()
            
            # 匹配"（一）"、"（二）"、"（三）"等章节标题
            if re.match(r"^[（(][一二三四五六七八九十百千万]+[）)]\s*", text):
                # 设置段落格式
                para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
                para.paragraph_format.left_indent = Pt(0)
                
                if original_text != text:
                    # 清空段落内容并重新设置
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = text
                    else:
                        para.add_run(text)
                    logger.info(f"修正章节标题缩进并清理空格: {text[:50]}")
                else:
                    logger.info(f"修正章节标题缩进: {text[:50]}")
    
    def _fix_references_indent(self):
        """
        取消参考文献章节下所有段落的首行缩进
        """
        import re
        in_references = False
        references_count = 0
        
        for para in self.doc.paragraphs:
            text = (para.text or "").strip()
            
            if "参考文献" in text and re.match(r"^[#\s]*参考文献", text):
                in_references = True
                logger.info(f"检测到参考文献章节: {text[:50]}")
                continue
            
            if in_references:
                if (re.match(r"^[（(][一二三四五六七八九十百千万]+[）)]", text) or 
                    re.match(r"^[0-9]+\.\s+", text) and not re.match(r"^[0-9]+\.[0-9]+", text)):
                    in_references = False
                    logger.info(f"离开参考文献章节，共处理 {references_count} 个段落")
                    continue
                
                # 在参考文献章节中，取消首行缩进
                if text:  # 非空段落
                    para.paragraph_format.first_line_indent = Pt(0)
                    references_count += 1
    
    def fill_from_markdown(self, markdown_input: str):
        import os
        if os.path.exists(markdown_input):
            logger.info(f"从文件填充: {markdown_input}")
            with open(markdown_input, "r", encoding="utf-8") as f:
                md_content = f.read()
        else:
            logger.info("从内容字符串填充")
            md_content = markdown_input
        # 将 &lt; 转换回 < 用于Word导出
        md_content = md_content.replace('&lt;', '<')
        parts = self._parse_markdown_to_parts(md_content)
        self._fill_by_parts(parts)
        self._fix_chapter_title_indent()
        self._fix_references_indent()

    def _setup_partial_protection(self):
        """
        设置部分保护：
        - 保留模板原有的保护设置
        - 只给新插入/修改的内容添加可编辑权限
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # 模板已经有正确的保护设置（readOnly + 段落级权限标记）
            # 我们只需要给新插入的段落添加可编辑权限
            
            added_count = 0
            perm_id = 10000  # 使用较大的起始ID避免与模板冲突
            
            for para in self.doc.paragraphs:
                text = (para.text or "").strip()
                if not text:
                    continue
                
                # 检查这段内容是否是我们新增的（不在模板原有内容中）
                is_new_content = text not in self.original_template_texts
                
                if is_new_content:
                    # 检查是否已经有编辑权限标记
                    has_perm = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}permStart') is not None
                    
                    if not has_perm:
                        # 添加可编辑权限
                        try:
                            perm_start = OxmlElement('w:permStart')
                            perm_start.set(qn('w:id'), str(perm_id))
                            perm_start.set(qn('w:edGrp'), 'everyone')
                            
                            perm_end = OxmlElement('w:permEnd')
                            perm_end.set(qn('w:id'), str(perm_id))
                            
                            para._element.addprevious(perm_start)
                            para._element.addnext(perm_end)
                            
                            added_count += 1
                            perm_id += 1
                        except Exception as e:
                            logger.debug(f"为段落添加编辑权限失败: {e}")
            
            logger.info(f"✓ 保护设置完成: 保留模板原有保护，为 {added_count} 个新段落添加编辑权限")
            
        except Exception as e:
            logger.warning(f"设置部分保护失败: {e}")
            import traceback
            logger.debug(traceback.format_exc())
    
    def save(self, output_path: str):
        self._setup_partial_protection()
        self.doc.save(output_path)

    def diagnose_template(self):
        logger.info("=" * 60)
        logger.info("Word 模板诊断报告")
        logger.info("=" * 60)

        titles = []
        for i, p in enumerate(self.doc.paragraphs):
            text = (p.text or "").strip()
            if not text:
                continue
            # 数字或中文数字开头 -> 认为是标题候选
            if re.match(r"^[（(]?[一二三四五六七八九十0-9]", text):
                titles.append((i, text))
                logger.info(f"段落 {i:3d}: {text[:80]}")

        logger.info("=" * 60)
        logger.info(f"共找到 {len(titles)} 个可能标题段落")
        logger.info("=" * 60)
        return titles


# ================== 示例用法 ==================
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    exporter = NSFCDocxExporter(
        template_path="/Users/yan/Developer/NOAHAI/NoahAgent/noah_agent/agent/nsfc/2026  面上项目-报告正文.docx"
    )
    exporter.fill_from_markdown("/Users/yan/Developer/NOAHAI/NoahAgent/noah_agent/outputs/nsfc_20251124_131745/国自然申请书.md")
    exporter.save("/Users/yan/Developer/NOAHAI/NoahAgent/noah_agent/outputs/nsfc_20251124_131745/filled_nsfc_application.docx")