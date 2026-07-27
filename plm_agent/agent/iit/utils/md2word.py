from typing import Literal

from docx import Document
from docx.shared import Inches, Pt
import requests


def md_to_word(
    input_file_path: str,
    output_file_path: str,
    logo_path: str,
    format_type: Literal["english", "chinese"] = "english",
) -> None:
    """将 Markdown 转为 Word；format_type 与转换服务约定一致：english / chinese 排版样式。"""
    try:
        # 使用 'rb' (二进制读取) 模式打开文件
        md_to_word_url = "https://test.noahai.co/markdown-to-word/convert"
        with open(input_file_path, 'rb') as f:
            # 构造 multipart/form-data 格式的数据
            # 注意：这里的字典键名 'file' 必须与你后端代码（如 Flask/FastAPI）中接收文件的字段名保持一致
            files = {'file': (input_file_path, f, 'text/markdown')}
            data = {'format_type': format_type}

            print(f"正在上传 {input_file_path} 并请求转换 (format_type={format_type})...")

            # 发送 POST 请求（与 curl -F format_type=... 等价）
            response = requests.post(md_to_word_url, files=files, data=data)

        # 4. 检查响应状态码
        if response.status_code == 200:
            # 如果成功，将返回的二进制流写入本地的 Word 文件
            with open(output_file_path, 'wb') as f_out:
                f_out.write(response.content)
            print(f"✅ 转换成功！文件已保存为: {output_file_path}")
            document = Document(str(output_file_path))
            # 中文排版由服务端生成，避免此处统一改为 Arial 破坏样式
            if format_type == "english":
                for paragraph in document.paragraphs:
                    for run in paragraph.runs:
                        # Detect if paragraph is a heading
                        is_heading = paragraph.style.name.startswith('Heading')

                        # Apply different font settings based on paragraph type
                        if is_heading:
                            run.font.name = 'Arial'
                            run.font.bold = True
                            # Keep existing heading size
                        else:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
            def apply_header_footer_word(doc):
                """Apply header and footer from a reference document if provided"""
                header = doc.sections[0].header
                # Set section margins
                # Set header margins
                header_paragraph = header.paragraphs[0]
                header_paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.CENTER (0=left, 1=center, 2=right)

                logo_run = header_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.96))
                # Lower the image by adjusting its vertical position
                # Add footer
                footer = doc.sections[0].footer
                footer_paragraph = footer.paragraphs[0]  # Use the first header_paragraph in the footer
                footer_paragraph.alignment = 1  # Center alignment
                footer_paragraph.text = "NOAH AI-AI Agent Specialized in Life Science"
                if logo_path == 'static/roche-logo.png':
                    footer_paragraph.text = "AI 智慧联  轻松做科研"
                footer_paragraph.style.font.size = Pt(9)  # Optional: adjust font size
            apply_header_footer_word(document)
            document.save(str(output_file_path))
            print(f"word已添加logo {output_file_path}")

        else:
            print(f"❌ 转换失败！状态码: {response.status_code}")
            print(f"错误信息: {response.text}")

    except FileNotFoundError:
        print(f"❌ 找不到文件: {input_file_path}，请确保该文件在当前目录下。")
    except requests.exceptions.RequestException as e:
        print(f"❌ 网络请求异常: {e}")
